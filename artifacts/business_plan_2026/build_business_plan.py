import json
import math
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path("/Users/jonasschroder/Desktop/ai_sports_prediction")
OUT = ROOT / "outputs/business_plan_2026"
ART = ROOT / "artifacts/business_plan_2026"
DATA = json.loads((OUT / "financial_summary.json").read_text())
OUT.mkdir(parents=True, exist_ok=True)

NAVY = "12263A"
BLUE = "1F5A7A"
CYAN = "DDEFF5"
PALE = "F3F7F9"
GREEN = "2E7D5B"
RED = "B54747"
AMBER = "B7791F"
GRAY = "68737D"
MID = "CBD5DC"
WHITE = "FFFFFF"
BLACK = "17212B"


def rgb(hex_color):
    return RGBColor.from_string(hex_color)


def set_cell_shading(cell, color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), color)


def set_cell_margins(cell, top=80, start=100, bottom=80, end=100):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def cant_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    tr_pr.append(OxmlElement("w:cantSplit"))


def repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_repeat_table_header(row):
    repeat_header(row)


def set_cell_width(cell, width_cm):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_cm * 567)))
    tc_w.set(qn("w:type"), "dxa")


def add_hyperlink(paragraph, text, url, color=BLUE, underline=True):
    part = paragraph.part
    rid = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rid)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    c = OxmlElement("w:color")
    c.set(qn("w:val"), color)
    r_pr.append(c)
    if underline:
        u = OxmlElement("w:u")
        u.set(qn("w:val"), "single")
        r_pr.append(u)
    new_run.append(r_pr)
    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


def add_field(paragraph, instruction):
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char)
    run._r.append(instr)
    run._r.append(fld_char2)


def fmt_eur(value, decimals=0):
    if value is None:
        return "–"
    s = f"{abs(value):,.{decimals}f}".replace(",", "X").replace(".", ",").replace("X", ".")
    return f"{'−' if value < 0 else ''}{s} €"


def fmt_num(value, decimals=0):
    return f"{value:,.{decimals}f}".replace(",", "X").replace(".", ",").replace("X", ".")


def fmt_pct(value, decimals=1):
    return f"{value * 100:.{decimals}f}%".replace(".", ",")


def fmt_te(value):
    rounded = int(round(value / 1000))
    return f"{'−' if rounded < 0 else ''}{abs(rounded)} T€"


def set_doc_defaults(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.82)
    section.bottom_margin = Inches(0.76)
    section.left_margin = Inches(0.88)
    section.right_margin = Inches(0.78)
    section.header_distance = Inches(0.3)
    section.footer_distance = Inches(0.3)
    section.different_first_page_header_footer = True

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = rgb(BLACK)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_after = Pt(7)
    normal.paragraph_format.line_spacing = 1.18

    for style_name, size, color, before, after in (
        ("Title", 30, NAVY, 0, 10),
        ("Subtitle", 14, BLUE, 0, 8),
        ("Heading 1", 17, BLUE, 18, 8),
        ("Heading 2", 13, BLUE, 12, 5),
        ("Heading 3", 11, NAVY, 9, 4),
    ):
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = rgb(color)
        style.font.bold = style_name != "Subtitle"
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    for name in ("List Bullet", "List Number"):
        styles[name].font.name = "Calibri"
        styles[name].font.size = Pt(10.5)
        styles[name].paragraph_format.left_indent = Cm(0.55)
        styles[name].paragraph_format.first_line_indent = Cm(-0.25)
        styles[name].paragraph_format.space_after = Pt(3)

    if "Small Note" not in styles:
        style = styles.add_style("Small Note", WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = "Calibri"
        style.font.size = Pt(8.5)
        style.font.color.rgb = rgb(GRAY)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.05
    if "Callout" not in styles:
        style = styles.add_style("Callout", WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = "Calibri"
        style.font.size = Pt(10.5)
        style.font.bold = True
        style.font.color.rgb = rgb(NAVY)
        style.paragraph_format.space_before = Pt(6)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.left_indent = Cm(0.35)
        style.paragraph_format.right_indent = Cm(0.35)

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = hp.add_run("AI SPORTS PREDICTION  ·  BUSINESSPLAN 2026")
    run.font.name = "Calibri"
    run.font.size = Pt(8)
    run.font.bold = True
    run.font.color.rgb = rgb(GRAY)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run("VERTRAULICH  ·  INTERNE PLANUNG  ·  ")
    run.font.size = Pt(8)
    run.font.color.rgb = rgb(GRAY)
    add_field(fp, "PAGE")


def add_para(doc, text="", style=None, bold_lead=None):
    p = doc.add_paragraph(style=style)
    if bold_lead and text.startswith(bold_lead):
        p.add_run(bold_lead).bold = True
        p.add_run(text[len(bold_lead):])
    else:
        p.add_run(text)
    return p


def add_bullets(doc, items):
    for item in items:
        add_para(doc, item, "List Bullet")


def add_numbered(doc, items):
    for item in items:
        add_para(doc, item, "List Number")


def add_callout(doc, title, text, color=CYAN):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, color)
    set_cell_margins(cell, 140, 180, 140, 180)
    p = cell.paragraphs[0]
    p.style = doc.styles["Callout"]
    p.add_run(title + " — ").bold = True
    p.add_run(text)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_table(doc, headers, rows, widths=None, font_size=8.5, header_color=BLUE, first_col_bold=True):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.style = "Table Grid"
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, h in enumerate(headers):
        c = hdr.cells[i]
        c.text = str(h)
        set_cell_shading(c, header_color)
        c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_margins(c)
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(0)
            for run in p.runs:
                run.font.name = "Calibri"
                run.font.size = Pt(font_size)
                run.font.bold = True
                run.font.color.rgb = rgb(WHITE)
        if widths:
            set_cell_width(c, widths[i])
    for ridx, row in enumerate(rows):
        cells = table.add_row().cells
        cant_split(table.rows[-1])
        for i, val in enumerate(row):
            c = cells[i]
            c.text = str(val)
            c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(c)
            if ridx % 2 == 1:
                set_cell_shading(c, PALE)
            if widths:
                set_cell_width(c, widths[i])
            for p in c.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT if i == 0 else WD_ALIGN_PARAGRAPH.RIGHT
                for run in p.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(font_size)
                    if i == 0 and first_col_bold:
                        run.font.bold = True
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def chapter(doc, number, title, new_page=True):
    if new_page:
        doc.add_page_break()
    doc.add_heading(f"{number}. {title}", level=1)


def add_source_ref(paragraph, refs):
    run = paragraph.add_run(" " + " ".join(f"[Q{x}]" for x in refs))
    run.font.size = Pt(8)
    run.font.color.rgb = rgb(BLUE)


def make_charts():
    width, height = 1500, 620
    margin = (120, 100, 70, 90)
    font_path = "/System/Library/Fonts/Supplemental/Arial.ttf"
    bold_path = "/System/Library/Fonts/Supplemental/Arial Bold.ttf"
    font = ImageFont.truetype(font_path, 24) if Path(font_path).exists() else ImageFont.load_default()
    small = ImageFont.truetype(font_path, 20) if Path(font_path).exists() else ImageFont.load_default()
    bold = ImageFont.truetype(bold_path, 30) if Path(bold_path).exists() else font
    colors = {"Pessimistisch": "#B54747", "Realistisch": "#1F5A7A", "Optimistisch": "#2E7D5B"}

    def frame(title, max_value):
        image = Image.new("RGB", (width, height), "white")
        draw = ImageDraw.Draw(image)
        draw.text((margin[0], 28), title, fill="#12263A", font=bold)
        left, top = margin[0], margin[1]
        right, bottom = width-margin[2], height-margin[3]
        for i in range(6):
            y = bottom - (bottom-top)*i/5
            draw.line((left, y, right, y), fill="#D8E0E5", width=2)
            draw.text((15, y-12), f"{max_value*i/5:.0f}", fill="#68737D", font=small)
        for i in range(5):
            x = left + (right-left)*i/4
            draw.text((x-30, bottom+18), f"Jahr {i+1}", fill="#68737D", font=small)
        return image, draw, (left, top, right, bottom)

    max_rev = max(y["revenue"]/1000 for d in DATA["scenarios"].values() for y in d["years"])
    image, draw, (left, top, right, bottom) = frame("Umsatzentwicklung nach Szenario (Tsd. €)", math.ceil(max_rev/100)*100)
    for idx, (name, data) in enumerate(DATA["scenarios"].items()):
        values = [y["revenue"]/1000 for y in data["years"]]
        points=[]
        for i,v in enumerate(values):
            x=left+(right-left)*i/4; y=bottom-(bottom-top)*v/(math.ceil(max_rev/100)*100); points.append((x,y))
        draw.line(points, fill=colors[name], width=6)
        for p in points: draw.ellipse((p[0]-7,p[1]-7,p[0]+7,p[1]+7),fill=colors[name])
        draw.text((left+idx*300, 70), name, fill=colors[name], font=font)
    p1 = ART / "chart_revenue.png"; image.save(p1)

    real = DATA["scenarios"]["Realistisch"]["years"]
    max_real = math.ceil(max(y["revenue"]/1000 for y in real)/50)*50
    image, draw, (left, top, right, bottom) = frame("Realistisches Szenario: Umsatz und EBITDA (Tsd. €)", max_real)
    group=(right-left)/5
    for i,yv in enumerate(real):
        center=left+group*(i+.5); bw=group*.24
        for j,(key,color) in enumerate((("revenue","#1F5A7A"),("ebitda","#72A8BA"))):
            value=yv[key]/1000; x0=center+(-1 if j==0 else 0.1)*bw; x1=x0+bw*.9
            y0=bottom-(bottom-top)*max(value,0)/max_real
            draw.rectangle((x0,y0,x1,bottom),fill=color)
    draw.text((left,70),"Umsatz",fill="#1F5A7A",font=font); draw.text((left+220,70),"EBITDA",fill="#72A8BA",font=font)
    p2 = ART / "chart_real.png"; image.save(p2)
    return p1, p2


chart_revenue, chart_real = make_charts()
doc = Document()
set_doc_defaults(doc)

# Cover
cover = doc.add_table(rows=1, cols=2)
cover.alignment = WD_TABLE_ALIGNMENT.CENTER
cover.autofit = False
cover.style = "Table Grid"
cover.cell(0, 0).text = ""
cover.cell(0, 1).text = ""
cover.cell(0, 0).width = Cm(4.0)
cover.cell(0, 1).width = Cm(12.5)
set_cell_shading(cover.cell(0, 0), NAVY)
set_cell_shading(cover.cell(0, 1), NAVY)
icon = ROOT / "apps/web/public/site-icon.png"
if icon.exists():
    cover.cell(0, 0).paragraphs[0].add_run().add_picture(str(icon), width=Cm(2.25))
    cover.cell(0, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
p = cover.cell(0, 1).paragraphs[0]
p.paragraph_format.space_after = Pt(5)
r = p.add_run("BUSINESSPLAN")
r.font.name = "Calibri"; r.font.size = Pt(13); r.font.bold = True; r.font.color.rgb = rgb("8FC6D4")
p2 = cover.cell(0, 1).add_paragraph()
r = p2.add_run("AI Sports Prediction")
r.font.name = "Calibri"; r.font.size = Pt(28); r.font.bold = True; r.font.color.rgb = rgb(WHITE)
p3 = cover.cell(0, 1).add_paragraph("Arbeitstitel · KI-gestützte Sportprognosen und Publisher-Widgets")
p3.runs[0].font.color.rgb = rgb(WHITE); p3.runs[0].font.size = Pt(12)
for c in cover.rows[0].cells:
    set_cell_margins(c, 520, 260, 520, 260)
    c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
r = p.add_run("Finale, nach Investment-Committee-Kritik überarbeitete Fassung")
r.font.size = Pt(13); r.font.bold = True; r.font.color.rgb = rgb(BLUE)
add_para(doc, "Planungszeitraum: August 2026 bis Juli 2031", "Subtitle")
add_para(doc, "Zweck: interne Unternehmensplanung sowie Gesprächsgrundlage für Steuerberatung, Versicherungen, Banken, Förderstellen und spätere Investoren.")
add_para(doc, "Gründer: Jonas Schröder und Jonas Schweisthal · je 50%, vorbehaltlich der IP- und Beteiligungsklärung mit der Universität.")

doc.add_paragraph().paragraph_format.space_after = Pt(30)
add_callout(doc, "Vertraulichkeit und Haftung", "Dieses Dokument ist eine interne Planungsunterlage und keine Rechts-, Steuer-, Anlage- oder Versicherungsberatung. Zahlen sind Modellrechnungen; Verträge, Steuern, Lizenzen und Deckungskonzepte sind von qualifizierten Beratern zu bestätigen.", PALE)
add_para(doc, "Stand: 20. Juli 2026 · geplanter Gründungsmonat: August 2026", "Small Note")

doc.add_page_break()
doc.add_heading("Dokumentenlogik und Evidenzklassen", level=1)
add_para(doc, "Der Plan trennt nachprüfbare Fakten, eigene Berechnungen und noch zu validierende Annahmen. Dadurch bleibt sichtbar, was bereits belegt ist und welche Aussagen vor externen Gesprächen nachgeschärft werden müssen.")
add_table(doc, ["Kennzeichnung", "Bedeutung", "Beispiel"], [
    ["Belegt", "Externe Primär-/Branchenquelle oder vorgelegter Screenshot", "69.919 Pageviews im dokumentierten Zeitraum"],
    ["Gründerangabe", "Vom Team berichtet, aber noch ohne prüffähiges Originaldokument", "Paper eingereicht und öffentlich archiviert; Identifier folgt"],
    ["Eigene Berechnung", "Rechenweg auf expliziten Inputs", "TAM/SAM/SOM und Finanzmodell"],
    ["Annahme", "Unsicherer Planungsparameter", "Churn, RPM, Conversion und CAC"],
], [3.0, 7.0, 7.2], 8.7)
add_callout(doc, "Zentrale Planungsentscheidung", "Das Widget-SaaS ist das Kernmodell. Werbeerlöse der kostenlosen Website bleiben ein ergänzender, volatiler Kanal. Eine externe Eigenkapitalrunde wird derzeit nicht empfohlen.")

doc.add_heading("Inhaltsübersicht", level=1)
contents = [
    "Executive Summary", "Unternehmen, Vision und langfristige Zielsetzung", "Problem und nachgewiesener Kundenbedarf",
    "Lösung sowie Produkt- und Leistungsbeschreibung", "Marktanalyse mit TAM, SAM und SOM", "Zielgruppen und Buyer Personas",
    "Wettbewerbsanalyse", "Positionierung und nachhaltige Wettbewerbsvorteile", "Geschäftsmodell, Preisgestaltung und Einnahmequellen",
    "Marketing-, Vertriebs- und Go-to-Market-Strategie", "Produktentwicklung, Technologie und geistiges Eigentum", "Operativer Plan, Prozesse und Ressourcen",
    "Gründerteam, Rollen und Personalbedarf", "Rechtsform, Beteiligungsstruktur und regulatorische Aspekte", "Meilenstein- und Umsetzungsplan 36 Monate",
    "Chancen, Risiken und Gegenmaßnahmen", "Finanzplanung über fünf Jahre", "Kapitalbedarf, Mittelverwendung und Finanzierungsszenarien",
    "Exit- bzw. langfristige Wachstumsstrategie", "Abschließende Investitionsargumentation", "Investment-Committee-Check", "Anhänge und Quellen"
]
toc_rows = []
for i in range(11):
    left = f"{i+1}. {contents[i]}"
    right = f"{i+12}. {contents[i+11]}" if i+11 < len(contents) else ""
    toc_rows.append([left, right])
add_table(doc, ["Kapitel 1–11", "Kapitel 12–Anhänge"], toc_rows, [8.6, 8.6], 8.2, first_col_bold=False)

# 1
chapter(doc, 1, "Executive Summary")
add_para(doc, "AI Sports Prediction ist ein geplantes deutsches Softwareunternehmen, das nachvollziehbare KI-Prognosen für Fußball, Tennis, NFL und NBA erzeugt und über eine kostenlose Consumer-Plattform sowie vier anpassbare B2B-Widgets ausspielt. Zielkunden sind Sport- und Nachrichtenredaktionen, digitale Publisher und spezialisierte Fanportale, die Prognoseformate ohne eigene Data-Science- und Widget-Infrastruktur veröffentlichen wollen.")
add_para(doc, "Die stärkste vorhandene Evidenz ist Reichweite, nicht Zahlungsbereitschaft: Die Vorgängerplattform LLM SoccerArena erzielte laut vorgelegten Google-Analytics-Screenshots vom 11. Juni bis 19. Juli 69.919 Pageviews, 36.066 Sitzungen, 13.669 aktive Nutzer und 13.596 neue Nutzer; die durchschnittliche Interaktionsdauer lag bei rund 2:22 Minuten. Ein Tagesschau-Video zum Projekt weist im Screenshot 7,7 Mio. Aufrufe aus. Diese Kennzahlen belegen Publikumsinteresse und PR-Fähigkeit, aber noch keine B2B-Retention oder zahlende Nachfrage.")
add_callout(doc, "Empfehlung", "Conditional Go: Gründung als schlanke UG (haftungsbeschränkt) ist vertretbar, wenn vor oder unmittelbar nach Gründung IP-, Daten- und Markenrechte geklärt und innerhalb von 90 Tagen mindestens 3–5 zahlende Designpartner gewonnen werden.")
add_table(doc, ["Dimension", "Bewertung", "Konsequenz"], [
    ["Produkt", "MVP nahezu marktreif; vier Widgettypen vorhanden/konzipiert", "Inhalte, Designs, Prompts, QA und Name finalisieren"],
    ["Markt", "Großer Digitalwerbemarkt; enges adressierbares Publisher-Segment", "Bottom-up verkaufen, nicht mit Marktgröße argumentieren"],
    ["Traktion", "Starke Event- und Medienreichweite", "Paid-Pilot-Conversion und Wiederkehrrate messen"],
    ["Finanzierung", "Kein institutionelles Kapital erforderlich", "15–25 T€ Liquiditätsreserve statt nur 2 T€"],
    ["Investierbarkeit", "Noch zu frühe Phase", "Nachweis von MRR, Churn, CAC und Prognosequalität erforderlich"],
], [3.2, 6.5, 7.4])

# 2
chapter(doc, 2, "Unternehmen, Vision und langfristige Zielsetzung")
add_para(doc, "Vision: AI Sports Prediction soll zur vertrauenswürdigen europäischen Referenz für verständliche, vorab veröffentlichte und nachträglich überprüfbare KI-Sportprognosen werden. Das Unternehmen verkauft nicht „sichere Tipps“, sondern eine publizierbare Informations- und Interaktionsschicht für Sportinhalte.")
add_para(doc, "Mission für die ersten drei Jahre: Publishern ermöglichen, innerhalb weniger Minuten ein gebrandetes, messbares Prognoseformat einzubinden, ohne selbst Datenpipelines, Modelle, redaktionelle Reasoning-Texte und Frontend-Komponenten entwickeln zu müssen.")
add_bullets(doc, [
    "12 Monate: 15 zahlende B2B-Kunden im realistischen Szenario, vier produktionsreife Widgettypen und belastbare Messung von Prognosegüte, Nutzung und Supportaufwand.",
    "24 Monate: 50 aktive B2B-Kunden, 250.000 monatliche Pageviews und wiederholbarer Founder-led-Sales-Prozess.",
    "36 Monate: 100 B2B-Kunden, dokumentierte Kohorten-Retention, erste bezahlte Rollen und Deutschland als funktionierender Kernmarkt.",
    "5 Jahre: 260 B2B-Kunden und 366,7 T€ Umsatz im realistischen Szenario; Internationalisierung nur nach Product-Market-Fit.",
])

# 3
chapter(doc, 3, "Problem und nachgewiesener Kundenbedarf")
add_para(doc, "Sportredaktionen benötigen regelmäßig interaktive, aktuelle und differenzierende Inhalte. Große Häuser können Daten- und Visualisierungsteams aufbauen; kleinere und mittlere Publisher müssen Datenquellen, Modellierung, Frontend, Hosting, Aktualisierung und redaktionelle Qualität mit begrenzten Ressourcen abdecken. Gleichzeitig stehen Publisher unter Druck, direkte Nutzung, Verweildauer und wiederkehrende Formate zu stärken. Der Reuters Institute Trends Report 2026 dokumentiert die strategische Unsicherheit durch rückläufige Suchzugriffe und höhere AI-Investitionen vieler Publisher. [Q9]")
add_para(doc, "Der Kundenbedarf ist bislang nur teilweise nachgewiesen. Medienreichweite und Website-Nutzung zeigen Interesse auf Nutzerseite. Interviews wurden nach Gründerangabe geführt, aber Anzahl, Rollen, Zahlungsbereitschaft und Transkripte liegen noch nicht als strukturierte Evidenz vor. Das wichtigste Risiko ist daher nicht die technische Machbarkeit, sondern die Frage, ob Redaktionen dauerhaft 49–499 € monatlich für das Format zahlen.")
add_table(doc, ["Hypothese", "Bestehende Evidenz", "Validierung bis Tag 90"], [
    ["Publisher wollen Prognosewidgets", "Interesse in Gesprächen; keine bezahlten Kunden", "30 strukturierte Interviews; 3–5 bezahlte Designpartner"],
    ["Reasoning steigert Nutzung", "2:22 min Engagement im Eventzeitraum; keine Widget-A/B-Tests", "A/B-Test Reasoning vs. Wahrscheinlichkeit"],
    ["Mehrere Sportarten erhöhen Wert", "Fußball-Prototyp; übrige Sportarten geplant", "Nur nach Kundenpriorität implementieren"],
    ["Automatisierter Outreach skaliert", "Agenten im Aufbau", "Antwort-, Termin- und Abschlussquote messen"],
], [5.1, 6.1, 5.9])

# 4
chapter(doc, 4, "Lösung sowie Produkt- und Leistungsbeschreibung")
add_para(doc, "Das Angebot besteht aus einer kostenlosen Plattform für Sportfans und einer B2B-Widget-/API-Schicht. Die Plattform schafft Reichweite, Datenfeedback und Markenbekanntheit; die Widgets monetarisieren redaktionelle Nutzung.")
add_table(doc, ["Widget", "Inhalt", "Primärer Einsatz"], [
    ["Prediction Card", "Prognose, Wahrscheinlichkeiten, Kernaussage und Reasoning", "Einzelspiel-/Matchartikel"],
    ["Match List", "Kommende Spiele mit kompakten Prognosen", "Liga-, Turnier- und Livecenter-Seiten"],
    ["Win Probability", "Visualisierte Siegwahrscheinlichkeiten", "Previews und Datengrafiken"],
    ["Key Factors", "Modellbegründung, Kontext und Unsicherheitsfaktoren", "Analyse- und Hintergrundstücke"],
], [3.7, 7.8, 5.7])
add_para(doc, "Kunden können Farben, Inhalte und Layout anpassen. Einbettung erfolgt per JavaScript/iframe bzw. API-gestützt; Verwaltung, Billing, Analytics und technischer Support werden zentral angeboten. Ausgaben müssen ein Prognosedatum, Datenstand, Modellversion, Unsicherheit und klare KI-Kennzeichnung enthalten.")

# 5
chapter(doc, 5, "Marktanalyse mit TAM, SAM und SOM")
add_para(doc, "Der relevante Markt ist kein allgemeiner „Sports-Tech-Markt“, sondern ein Schnittsegment aus Publisher-Software, interaktiven Sportinhalten und digitaler Monetarisierung. IAB Europe beziffert den europäischen digitalen Werbemarkt 2025 auf 131 Mrd. €; der deutsche Markt für digitale Display- und Videowerbung soll 2026 laut OVK/BVDW auf rund 8,2 Mrd. € wachsen. In den USA erreichten digitale Werbeerlöse 2025 nach IAB 294,6 Mrd. US-Dollar. Diese Zahlen zeigen ein großes Erlösökosystem, sind jedoch nicht der adressierbare Umsatz der Gesellschaft. [Q3] [Q4] [Q5]")
add_para(doc, "Eurostat weist für 2023 vorläufig 76.328 Publishing-Unternehmen in der EU und 5.374 in Deutschland aus. Daraus wird ein transparenter, bottom-up Marktproxy abgeleitet; weder Sportrelevanz noch Zahlungsbereitschaft sind in der Statistik enthalten. [Q6]")
add_table(doc, ["Ebene", "Rechnung", "Bandbreite p.a.", "Interpretation"], [
    ["TAM EU B2B", "76.328 × 5–15% sportrelevant × 1–3 T€ ACV", "3,8–34,3 Mio. €", "Eigene Berechnung; sehr breite Obergrenze"],
    ["SAM Deutschland", "5.374 × 5–15% × 1–2 T€ ACV", "0,27–1,61 Mio. €", "Deutschland-first; Publisher-Proxy"],
    ["SOM 36 Monate", "50–150 Kunden × 1,2–1,8 T€ ACV", "0,06–0,27 Mio. €", "Operativ erreichbare Bandbreite"],
], [3.0, 6.7, 3.5, 5.0])
add_para(doc, "Die SAM-Bandbreite ist klein genug, dass Deutschland allein kein sehr großes Venture-Case rechtfertigt. Ein skalierbarer internationaler Case verlangt entweder höhere Enterprise-ARPA, zusätzliche Lizenz-/Datenprodukte oder Expansion in weitere Märkte. Das realistische Fünfjahresmodell bleibt deshalb bewusst unter 0,4 Mio. € Umsatz.")

# 6
chapter(doc, 6, "Zielgruppen und konkrete Buyer Personas")
add_table(doc, ["Persona", "Ziel / Schmerz", "Kaufkriterium", "Einwand", "Passendes Paket"], [
    ["Sportredakteur/in eines Regionalverlags", "Mehr aktuelle Sportinhalte ohne Entwicklerteam", "Schnelle Einbindung, journalistisch nutzbarer Text", "„Ist das verlässlich und redaktionell vertretbar?“", "Starter/Growth"],
    ["Head of Product / Audience", "Engagement und wiederkehrende Nutzung", "Analytics, Branding, Performance, Datenschutz", "„Welchen Uplift erzeugt das Widget?“", "Growth/Enterprise"],
    ["Betreiber eines Sportportals", "Breite Spielabdeckung und Automatisierung", "API-Stabilität, Sportarten, Kosten", "„Warum nicht kostenlose Widgets?“", "Growth"],
    ["Digitalagentur für Publisher", "Wiederholbare Lösung für mehrere Kunden", "Mandantenfähigkeit, SLA, White Label", "„Können wir weiterverkaufen?“", "Enterprise"],
], [3.6, 4.5, 4.6, 4.3, 2.3], 7.7)
add_para(doc, "Priorität 1 sind kleine und mittlere deutsche Sport-/Nachrichtenredaktionen mit klarer digitaler Verantwortung und ohne eigene ML-Produktteams. Nicht priorisiert werden Wettanbieter, reine Affiliate-Seiten und US-Medien, solange Rechte, englische Inhalte und Support nicht ausreichend abgesichert sind.")

# 7
chapter(doc, 7, "Wettbewerbsanalyse einschließlich direkter und indirekter Wettbewerber")
add_table(doc, ["Anbieter/Alternative", "Stärke", "Lücke/Abgrenzung", "Implikation"], [
    ["Opta / Stats Perform", "Datenbreite, Marke, etablierte Modelle", "Premium-Positionierung; nicht primär schlankes SME-Widget", "Nicht auf Datenbreite konkurrieren"],
    ["Sportmonks Widgets", "Fertige Fußballwidgets; Preise ab 59 €/Monat bei Jahreszahlung", "Fokus auf Daten/Widgets, nicht zwingend erklärbare KI-Prognosen", "Starterpreis ist plausibel; Nutzen muss klarer sein"],
    ["API-Sports", "Multisport-APIs und anpassbare Widgets", "Kunde trägt mehr Produkt-/Redaktionsarbeit", "Einfachheit und Reasoning verkaufen"],
    ["ScoreAxis", "Kostenlose Widgets", "Geringere Differenzierung/Service", "Kostenlose Alternativen machen reine Einbettung austauschbar"],
    ["Buchmacherquoten", "Starke Marktaggregation und Aktualität", "Wettkontext, keine neutrale redaktionelle Erklärung", "Kalibrierung benchmarken, aber nicht kopieren"],
    ["Interne Data-Science-Teams", "Kontrolle, proprietäre Daten", "Hohe Fixkosten und lange Time-to-Market", "Auf kleinere/mittlere Häuser fokussieren"],
], [3.7, 4.9, 5.4, 4.1], 7.4)
add_para(doc, "Opta beschreibt seine Supercomputer-Prognosen als Kombination aus Wettquoten und Power Rankings. Das ist ein relevanter Benchmark, aber kein Beleg für Gleichwertigkeit des eigenen Systems. [Q11] Sportmonks, API-Sports und ScoreAxis zeigen, dass Widgets technisch und preislich bereits verfügbar sind. [Q12] [Q13] [Q14]")

# 8
chapter(doc, 8, "Positionierung und nachhaltige Wettbewerbsvorteile")
add_callout(doc, "Positionierungssatz", "Für digitale Sportredaktionen ohne eigenes ML-Produktteam liefert AI Sports Prediction gebrandete, erklärbare Prognosewidgets, die in wenigen Minuten integrierbar, vorab versioniert und nach dem Event messbar sind.")
add_para(doc, "Prompts allein sind kein belastbarer Schutzwall: Sie können rekonstruiert oder durch Modellwechsel entwertet werden. Nachhaltiger sind vier miteinander verknüpfte Assets:")
add_bullets(doc, [
    "eine lückenlose, vor Spielbeginn zeitgestempelte Prognosehistorie mit Kalibrierung und Fehleranalyse;",
    "publisher-spezifische Integrationen, Analytics, Branding und Workflow-Einbettung;",
    "ein wissenschaftlich nachvollziehbarer Benchmark gegen naive Baselines, Marktquoten und etablierte Systeme;",
    "Distribution, Kundenbeziehungen und eine vertrauenswürdige Marke für transparentes Reasoning.",
])
add_para(doc, "Das eingereichte/archivierte Paper ist nach Gründerangabe relevant; bis URL, Version, Reviewstatus und konkrete Ergebnisse vorliegen, darf der Plan keine Gleichwertigkeit mit Opta oder Buchmachern behaupten. Zulässige Formulierung: „Die Methode wird wissenschaftlich evaluiert; belastbare Vergleichsergebnisse werden nach Veröffentlichung verlinkt.“")

# 9
chapter(doc, 9, "Geschäftsmodell, Preisgestaltung und Einnahmequellen")
add_table(doc, ["Paket", "Monatlich", "Erstes Vertragsjahr jährlich", "Zielkunde", "Arbeitsannahme Leistungsgrenze"], [
    ["Starter", "49 €", "539 € (= 11 × 49 €)", "kleine Redaktion/Portal", "1 Domain, Basisbranding, Standard-Support"],
    ["Growth", "149 €", "1.639 € (= 11 × 149 €)", "wachsende Redaktion", "mehr Domains/Traffic, Analytics, Priorität"],
    ["Enterprise", "individuell; Modell 499 €/Monat", "individuell", "Verlag/Agentur", "SLA, White Label, Volumen, individuelle Rechte"],
], [2.8, 2.7, 4.4, 4.0, 6.0], 7.8)
add_para(doc, "Vertragslogik: 12 Monate Mindestlaufzeit; danach monatliche Verlängerung, soweit AGB und B2B-Vertragsrecht dies wirksam abbilden. Die jährliche Vorauszahlung gewährt ungefähr einen Gratismonat und verbessert Liquidität. Preisgrenzen dürfen nicht nur auf Pageviews beruhen, sondern sollten Domains, eingebundene Artikel, Impressions, Sportarten, API-Aufrufe, Support und SLA kombinieren.")
add_para(doc, "Werbung auf der kostenlosen Website wird aus Pageviews × Netto-RPM modelliert. Google veröffentlicht Revenue-Share-Informationen, garantiert aber keinen RPM. Daher werden 2,50 €, 5 € und 8 € netto je 1.000 Pageviews als explizite Szenarioannahmen verwendet. [Q15]")
add_callout(doc, "Strategische Regel", "Keine Wettaffiliate-Erlöse in der Startphase. Sie würden das Informationsversprechen, die Versicherbarkeit und die regulatorische Position unnötig belasten.")

# 10
chapter(doc, 10, "Marketing-, Vertriebs- und Go-to-Market-Strategie")
add_para(doc, "Der Markteintritt erfolgt Deutschland-first über Founder-led Sales. Die vorhandene Medienresonanz dient als Gesprächsöffner, nicht als Ersatz für einen Kaufgrund. Jeder Pilot braucht einen Sponsor, eine definierte Einbindung, Erfolgsmetriken und einen Preis.")
add_table(doc, ["Phase", "Ziel", "Maßnahmen", "KPI / Gate"], [
    ["0–30 Tage", "Problem und Angebot schärfen", "30 Zielaccounts, 15 Interviews, Demo, 4 Widgetdesigns", "≥5 konkrete Pilotinteressen"],
    ["31–60 Tage", "Paid Design Partners", "Individuelle Einbindung, 30-Tage-Test gegen 49–149 €", "≥3 zahlende Kunden"],
    ["61–90 Tage", "Wiederholbarkeit", "Case Study, Referral, standardisiertes Onboarding", "≥5 zahlende Kunden; <8 h Onboarding"],
    ["Monat 4–12", "Skalieren", "100–150 qualifizierte Kontakte/Monat, Branchenpartnerschaften", "15 Kunden; Pipeline-Coverage ≥3×"],
], [2.6, 3.4, 7.7, 4.6])
add_para(doc, "Outreach-Agenten dürfen Recherche, Personalisierung und Monitoring unterstützen. Vor automatischem Versand sind DSGVO/UWG, Datenherkunft, Opt-out, Frequenz und menschliche Freigabe zu prüfen. Erfolgsmetriken: Zustellrate, positive Antwortquote, Terminquote, Pilotquote, Abschlussquote, Cash-CAC und Zeit-CAC.")
add_bullets(doc, [
    "Content: wöchentlicher transparenter Forecast-Scorecard-Artikel statt selektiver Trefferkommunikation.",
    "PR: wissenschaftliche Methodik und Messbarkeit; keine Superlative ohne Benchmark.",
    "Sales: sportliches Ereigniskalender-Playbook 8–12 Wochen vor Turnieren/Saisonstarts.",
    "Retention: monatlicher Publisher-Report mit Impressions, CTR/Engagement, Ladezeit und Nutzung pro Widget.",
])

# 11
chapter(doc, 11, "Produktentwicklung, Technologie und geistiges Eigentum")
add_para(doc, "Die bestehende Codebasis umfasst Webanwendung, API, Worker, Datenbank, Widget-Konfiguration, Billing-/Kundenportal, Analytics sowie Ansätze für Outreach und Monitoring. Vor Marktstart sind Produktreife und Nachweisführung wichtiger als weitere Featurebreite.")
add_table(doc, ["Priorität", "Produkt-/Technikarbeit", "Abnahmekriterium"], [
    ["P0", "Name/Domain, vier Widgetdesigns, Prognoseinhalte, Prompt-/Modellversionierung", "Produktionsdemo auf 3 Testdomains"],
    ["P0", "Monitoring, Caching, Fehlerzustände, Rollback, Datenstand", "Alarmierung + dokumentierter Incident-Prozess"],
    ["P0", "Timestamped Forecast Ledger und Ergebnisimport", "Kein rückwirkendes Überschreiben; reproduzierbare Scorecard"],
    ["P1", "Publisher Analytics und Self-Service Branding", "Monatlicher Nutzungsreport"],
    ["P1", "Tennis/NFL/NBA", "Nur nach 2+ zahlenden Kundenanforderungen je Sport"],
], [2.2, 9.0, 7.1])
add_para(doc, "IP-Gate: Vor Gründung müssen beide Gründer dokumentieren, wann und unter welchen Arbeits-/Forschungsbedingungen Code, Modelle, Prompts, Datensätze und Paper entstanden. Mit der LMU-Transferstelle sind Erfindungsmeldung, Softwareurheberrecht, Nebentätigkeit, Forschungsdaten und mögliche Beteiligungsansprüche zu klären. [Q18]")
add_para(doc, "Daten-/Logo-Gate: Auch wenn Teamlogos nur zur Identifikation dienen, können Urheber-, Marken- und Datenbankrechte betroffen sein. Für TheSportsDB und jede weitere API ist eine schriftlich archivierte kommerzielle Lizenzprüfung erforderlich. Fehlt sie, sind neutrale Text-/Abkürzungsdarstellungen und lizenzierte Ersatzassets zu verwenden.")

# 12
chapter(doc, 12, "Operativer Plan, Prozesse und benötigte Ressourcen")
add_table(doc, ["Prozess", "Owner", "Rhythmus", "Kontrolle"], [
    ["Datenimport & Prognosen", "Founder Tech", "ereignisabhängig/täglich", "Freshness, Vollständigkeit, Drift"],
    ["Editorial QA", "wechselnd", "vor Veröffentlichung", "Unsicherheit, Quellen, unzulässige Aussagen"],
    ["Widgetbetrieb", "Founder Tech", "laufend", "Uptime, Ladezeit, Fehlerquote"],
    ["Sales & Onboarding", "Founder Commercial", "wöchentlich", "Pipeline, Conversion, Onboardingstunden"],
    ["Billing & Support", "Founder Commercial", "monatlich/laufend", "DSO, Tickets, SLA"],
    ["Security & Backups", "Founder Tech", "wöchentlich/monatlich", "Restore-Test, Secrets, Zugriff"],
], [4.0, 3.4, 3.4, 7.4])
add_para(doc, "Aktuelle Infrastrukturkosten betragen nach Gründerangabe rund 80 € AWS und 10 € Sportdaten-API pro Monat. Prognosen werden vorab berechnet und gecacht; zusätzliche Widgetaufrufe lösen innerhalb der vorhandenen Kontingente praktisch keine zusätzlichen Rechen- oder Datenkosten aus. AWS, Sportdaten-API, LLM-Betrieb und Founder-Support werden deshalb als fixe beziehungsweise stufenfixe Kosten eingeordnet. Das Modell ergänzt Software/Admin, Buchhaltung, Versicherung, Recht/Compliance und Marketing. Anbieterlimits und der Punkt des nächsten Kostensprungs sind durch Lasttests zu dokumentieren.")
add_bullets(doc, [
    "Serviceziel zum Start: 99,5% monatliche Verfügbarkeit ohne hartes SLA; Enterprise-SLA erst nach Messhistorie.",
    "Recovery: tägliche Datenbanksicherung, monatlicher Restore-Test, dokumentierte RPO/RTO-Ziele.",
    "Support: klare Geschäftszeiten, Severity-Stufen und Statusseite; keine pauschale 24/7-Zusage.",
    "Vendor Risk: mindestens eine alternative Sportdatenquelle und austauschbare Modellprovider-Schnittstelle.",
])

# 13
chapter(doc, 13, "Gründerteam, Rollen und zukünftiger Personalbedarf")
add_para(doc, "Jonas Schröder und Jonas Schweisthal forschen nach Gründerangabe am Munich Center for Machine Learning/LMU im Bereich Causal Machine Learning und promovieren. Beide sind zu gleichen Teilen als Founder vorgesehen und können jeweils ungefähr zehn Stunden pro Woche beitragen. Die fachliche Stärke liegt klar in ML, Evaluation und Software; die größte Kompetenzlücke liegt in B2B-Vertrieb, Publisher-Produkten und kommerzieller Verantwortung.")
add_table(doc, ["Zeitraum", "Rollenmodell", "Cash-Personalaufwand", "Auslöser"], [
    ["M1–M24", "2 Founder, je ca. 10 h/Woche", "0 € Gehalt im Modell", "Neben PhD; Schattenkosten 5.196 €/Monat"],
    ["Jahr 3", "Teilzeit Sales/Customer Success", "30 T€ realistisch", ">75 Ø Kunden/definierte Pipeline"],
    ["Jahr 4", "Sales/CS + Engineering/Operations", "132 T€ realistisch", "170 Kunden, Support-/Uptimebedarf"],
    ["Jahr 5", "kleines Kernteam", "190 T€ realistisch", "260 Kunden, Internationalisierungsvorbereitung"],
], [3.0, 6.0, 4.2, 5.0])
add_para(doc, "Die 50/50-Struktur braucht einen Gesellschaftervertrag mit Deadlock-Verfahren, Reserved Matters, Vesting/Reverse Vesting, Good-/Bad-Leaver, IP-Übertragung, Zeitcommitment, Nebenbeschäftigung, Wettbewerbsregeln und Konfliktlösung. Gleichvergütung ist nur sinnvoll, solange Beiträge und Verantwortungen vergleichbar sind; Abweichungen müssen transparent beschlossen werden.")

# 14
chapter(doc, 14, "Rechtsform, Beteiligungsstruktur und regulatorische Aspekte")
add_para(doc, "Empfehlung: UG (haftungsbeschränkt) mit 2.000 € Stammkapital, ergänzt um eine dokumentierte Gesellschafterdarlehens-/Liquiditätszusage. Eine GmbH verlangt 25.000 € Stammkapital; bei Bargründung müssen vor Anmeldung grundsätzlich mindestens 12.500 € eingezahlt sein. Mit 2.000 € ist die UG die realistische haftungsbeschränkte Startform. [Q19]")
add_table(doc, ["Thema", "Erforderliche Maßnahme", "Status/Gate"], [
    ["Sitz München vs. Grünwald", "Tatsächlichen Ort der Geschäftsleitung, Betriebsstätte und Gewerbesteuer prüfen", "München konservativ modelliert; keine Briefkastengestaltung"],
    ["AI Act", "KI-generierte Inhalte und Systemgrenzen transparent kennzeichnen", "Transparenzpflichten ab 02.08.2026 prüfen [Q16]"],
    ["Datenschutz/TDDDG", "CMP, Cookie-/Ad-Tech-Prüfung, AVV/DPA, Löschkonzept, Verzeichnis", "Vor Tracking/Ads live"],
    ["B2B-Verträge", "Leistungsumfang, SLA, Vergütung, Laufzeit, IP, AVV, Haftung, Support", "Anwaltlich prüfen"],
    ["E-Rechnung/USt", "Empfangs-/Ausgangsprozess, Reverse Charge bei Ausland, USt-ID", "Mit Steuerberatung einrichten [Q20]"],
    ["Daten/Logos", "Kommerzielle Nutzungsrechte schriftlich sichern", "Launch-Gate"],
], [4.1, 8.7, 5.3])
doc.add_heading("Haftung, Prognosen und Sportwetten", level=2)
add_para(doc, "Website und Widgets stellen statistische Informationen und Unterhaltung bereit, keine Wetteinsätze, Wettvermittlung oder Erfolgsgarantie. Startseitig und je Prognose sind Unsicherheit, Datenstand und Modellcharakter kenntlich zu machen. Ein Hinweis „keine Haftung“ beseitigt gesetzliche Haftung nicht. Insbesondere lassen sich Vorsatz, grobe Fahrlässigkeit, Verletzung von Leben/Körper/Gesundheit und zwingende Haftung nicht pauschal ausschließen. Haftungsbegrenzungen müssen adressatengerecht in B2B-AGB und Verträgen formuliert werden.")
add_para(doc, "Empfohlene Deckungen: IT-/Vermögensschadenhaftpflicht mit Medien-, IP-, Datenschutz- und Eigenschadenbausteinen; Cyberversicherung einschließlich Incident Response; Betriebshaftpflicht; später D&O/Managerhaftpflicht und Firmenrechtsschutz. Prämien werden nicht geschätzt, sondern anhand eines vollständigen Risikofragebogens eingeholt. Die IHK betont die bedarfsbezogene Prüfung betrieblicher Versicherungen. [Q21]")

# 15
chapter(doc, 15, "Meilenstein- und Umsetzungsplan für die nächsten 36 Monate")
add_table(doc, ["Zeitraum", "Meilenstein", "Messbares Ergebnis", "Abbruch-/Pivot-Gate"], [
    ["Aug–Okt 2026", "Gründung & Launch", "UG, Verträge, Rechte, 4 Widgets, 3 zahlende Partner", "Ohne Rechtefreigabe kein öffentlicher Launch"],
    ["Nov 2026–Jan 2027", "Paid-Pilot-Beweis", "5 Kunden, Onboarding <8 h, erste Case Study", "<2 zahlende Kunden: Segment/Preis neu testen"],
    ["Feb–Jul 2027", "Founder-led Sales", "15 Kunden, 75k PV/Monat, Scorecard", "Churn >3%/Monat: Retention vor Wachstum"],
    ["Aug 2027–Jan 2028", "Produktisierung", "30 Kunden, Self-Service-Billing, Analytics", "Support >2 h/Kunde/Monat: Produkt vereinfachen"],
    ["Feb–Jul 2028", "Wiederholbarkeit", "50 Kunden, 250k PV/Monat, positive Cash-Run-rate", "CAC-Payback >12 Monate: Kanal stoppen"],
    ["Aug 2028–Jan 2029", "Team-Gate", "75 Kunden, erste Teilzeitrolle vorbereitet", "Keine Einstellung ohne 12 Monate Runway"],
    ["Feb–Jul 2029", "Deutschland-PMF", "100 Kunden, dokumentierte Retention und Referenzen", "Internationalisierung erst bei NRR ≥100%"],
], [3.3, 4.2, 6.8, 5.5], 7.7)

# 16
chapter(doc, 16, "Chancen, Risiken und konkrete Gegenmaßnahmen")
add_table(doc, ["Risiko", "Wahrscheinlichkeit", "Auswirkung", "Gegenmaßnahme", "Frühindikator"], [
    ["Keine Zahlungsbereitschaft", "hoch", "sehr hoch", "bezahlte Designpartner, Preisinterviews, enger ICP", "Pilot→Paid <30%"],
    ["Eventtraffic fällt ab", "hoch", "mittel", "SaaS zuerst; Always-on Sportkalender/SEO/Newsletter", "PV außerhalb Events"],
    ["Daten-/Logorechte unklar", "mittel-hoch", "sehr hoch", "schriftliche Lizenz, Ersatzassets, Rechtsprüfung", "fehlende Vertragsklausel"],
    ["Prognosegüte/Vertrauen", "mittel", "hoch", "öffentliche Kalibrierung, Baselines, Versionierung", "Brier/Log Loss driftet"],
    ["Founder-Kapazität", "hoch", "hoch", "enge Roadmap, Rotationsplan, Automatisierung mit Kontrolle", "Support-/Sales-SLAs verfehlt"],
    ["50/50-Deadlock", "mittel", "hoch", "Gesellschaftervertrag, Mediations-/Buy-sell-Mechanismus", "wiederholte Blockaden"],
    ["Vendor-/Cloud-Ausfall", "mittel", "mittel", "Caching, Backups, Fallback-Provider, Statusseite", "Fehlerrate/Latenz"],
    ["Haftung/Datenschutz", "mittel", "hoch", "AGB, DPA, CMP, Security, Versicherungen", "Incidents/Beschwerden"],
], [3.5, 2.6, 2.5, 6.9, 4.4], 7.3)
add_para(doc, "Größte Chance ist die Kombination aus glaubwürdiger ML-Kompetenz, vorhandener Reichweitenresonanz und einem klaren, kleinen Publisher-Problem. Größtes Risiko ist, diese Resonanz fälschlich als Product-Market-Fit zu behandeln.")

# 17
chapter(doc, 17, "Finanzplanung für mindestens fünf Jahre")
add_para(doc, "Das Modell plant die ersten 24 Monate monatlich und die Jahre 3–5 jährlich. Beträge sind netto. Gründergehälter betragen in M1–M24 cash-seitig 0 €, die Opportunitätskosten von 5.196 € pro Monat werden separat als Schattenrechnung gezeigt. Die Steuerbasis ist konservativ München; Verlustvorträge werden berücksichtigt. Das vollständige, formelgestützte Modell wird als Excel-Datei mitgeliefert.")
add_heading = doc.add_heading
doc.add_heading("Kernannahmen", level=2)
add_table(doc, ["Treiber", "Pessimistisch", "Realistisch", "Optimistisch"], [
    ["B2B-Kunden Ende M12 / M24 / J5", "5 / 15 / 45", "15 / 50 / 260", "30 / 100 / 650"],
    ["Pageviews M12 / M24 / J5", "25k / 60k / 200k", "75k / 250k / 1,5 Mio.", "150k / 700k / 5 Mio."],
    ["Netto-RPM", "2,50 €", "5,00 €", "8,00 €"],
    ["Monatlicher Churn nach Mindestlaufzeit", "3,0%", "2,0%", "1,0%"],
    ["Variable Daten-/LLM-/Supportkosten", "0%", "0%", "0%"],
    ["Kostenklassifikation", "API/Cloud stufenfix", "API/Cloud stufenfix", "API/Cloud stufenfix"],
    ["Vollkosten-CAC", "1.500 €", "750 €", "350 €"],
], [5.1, 4.0, 4.0, 4.0])

doc.add_heading("Fünfjahresergebnis", level=2)
for scen_name in ("Pessimistisch", "Realistisch", "Optimistisch"):
    if scen_name != "Pessimistisch":
        doc.add_page_break()
    years = DATA["scenarios"][scen_name]["years"]
    rows = []
    for label, key in (("Umsatz", "revenue"), ("Variable Kosten", "variable"), ("Fixkosten", "fixed"), ("Personalaufwand", "payroll"), ("EBITDA", "ebitda"), ("EBIT", "ebit"), ("Steuern", "tax"), ("Jahresergebnis", "netIncome"), ("Netto-Cashflow", "netCf"), ("Endliquidität", "endCash")):
        rows.append([label] + [fmt_eur(y[key]) for y in years])
    doc.add_heading(scen_name, level=3)
    add_table(doc, ["Kennzahl", "Jahr 1", "Jahr 2", "Jahr 3", "Jahr 4", "Jahr 5"], rows, [4.2, 2.55, 2.55, 2.55, 2.55, 2.55], 7.2)

doc.add_picture(str(chart_revenue), width=Cm(16.8))
p = doc.paragraphs[-1]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_para(doc, "Abbildung: Eigene Szenariorechnung. Die große Bandbreite unterstreicht, dass Kundenakquise, ARPA und Traffic – nicht Infrastrukturkosten – die Ergebnisunsicherheit dominieren.", "Small Note")

doc.add_page_break()
doc.add_heading("Realistische Monatsplanung M1–M24", level=2)
real_months = DATA["scenarios"]["Realistisch"]["months"]
metric_specs = [
    ("Aktive Kunden", "active", lambda v: fmt_num(v, 1)),
    ("Pageviews", "pageviews", lambda v: fmt_num(v, 0)),
    ("Umsatz", "revenue", fmt_eur),
    ("Variable Kosten", "variable", fmt_eur),
    ("Fix-/Einmalkosten", None, None),
    ("Personal cash", None, None),
    ("EBITDA", "ebitda", fmt_eur),
    ("EBIT", "ebit", fmt_eur),
    ("Steuern", "tax", fmt_eur),
    ("Jahresergebnis", "netIncome", fmt_eur),
    ("Netto-Cashflow", "netCf", fmt_eur),
    ("Endliquidität", "endCash", fmt_eur),
]
for block_start in range(0, 24, 6):
    if block_start == 12:
        doc.add_page_break()
    block = real_months[block_start:block_start+6]
    headers = ["Kennzahl"] + [f"M{x['m']}\n{x['date'][5:7]}/{x['date'][2:4]}" for x in block]
    rows = []
    for label, key, formatter in metric_specs:
        values = []
        for m in block:
            if label == "Fix-/Einmalkosten":
                values.append(fmt_eur(m["fixed"] + m["oneTime"]))
            elif label == "Personal cash":
                values.append("0 €")
            else:
                values.append(formatter(m[key]))
        rows.append([label] + values)
    add_table(doc, headers, rows, [4.3] + [2.05]*6, 6.8)

doc.add_page_break()
doc.add_heading("Break-even und Unit Economics", level=2)
unit_rows=[]
for label, key, kind in (
    ("Blended MRR", "mrr", "eur"), ("Deckungsbeitragsmarge", "gm", "pct"), ("LTV", "ltv", "eur"),
    ("Cash-CAC", "cashCac", "eur"), ("Vollkosten-CAC", "fullCac", "eur"), ("LTV/Vollkosten-CAC", "ltvFull", "x"),
    ("Vollkosten-Payback", "fullPayback", "months"), ("Cash Break-even Kunden ohne Ads", "breakEvenCustomers", "count")
):
    vals=[]
    for scen in ("Pessimistisch", "Realistisch", "Optimistisch"):
        v=DATA["scenarios"][scen]["unit"][key]
        vals.append(fmt_eur(v,1) if kind=="eur" else fmt_pct(v) if kind=="pct" else f"{v:.1f}×".replace(".",",") if kind=="x" else f"{v:.1f} Monate".replace(".",",") if kind=="months" else fmt_num(v,1))
    unit_rows.append([label]+vals)
add_table(doc,["Kennzahl","Pessimistisch","Realistisch","Optimistisch"],unit_rows,[5.8,4.0,4.0,4.0],8.0)
real_unit = DATA["scenarios"]["Realistisch"]["unit"]
add_para(doc, f"Im realistischen Modell beträgt der Blended MRR {fmt_eur(real_unit['mrr'], 2)}, die modellierte Deckungsbeitragsmarge {fmt_pct(real_unit['gm'])}, der vereinfachte LTV {fmt_eur(real_unit['ltv'], 2)} und das LTV/Vollkosten-CAC-Verhältnis {str(round(real_unit['ltvFull'], 2)).replace('.', ',')}×. Die Marge von 100% bedeutet nicht, dass der Betrieb kostenlos ist: Die vorhandenen AWS-, API-, LLM- und Supportressourcen stehen in den Fix- und Personalkosten. Diese Unit Economics sind Planannahmen und gelten nur bis zum nächsten Kapazitäts- oder Lizenzsprung. Besonders der LTV reagiert stark auf Churn und ist erst nach mindestens 12–18 Monaten Kohortendaten belastbar.")
doc.add_picture(str(chart_real), width=Cm(16.8)); doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER

doc.add_page_break()
doc.add_heading("Sensitivitätsanalyse", level=2)
base_y5=DATA["scenarios"]["Realistisch"]["years"][4]
add_para(doc, "Bei unveränderten Werbeerlösen, Fix- und Personalkosten liegt das Y5-EBITDA im realistischen Szenario bei 50% des geplanten Kundenvolumens und 80% ARPA bei rund −40 T€. Bei 50% Kundenvolumen und 110% ARPA wird der Break-even knapp überschritten; bei 75% Kundenvolumen und 80% ARPA liegt das EBITDA bereits bei rund 19 T€. Kundenzahl und ARPA bleiben damit die dominanten Hebel. Ein Werbe-RPM von 2,50 € statt 5 € reduziert den Y5-Werbeumsatz bei 1,5 Mio. monatlichen Pageviews um 36 T€ pro Jahr.")
arps = [.8, .9, 1, 1.1, 1.2]
custs = [.5, .75, 1, 1.25, 1.5]
sens_rows = []
for cm in custs:
    values = []
    for am in arps:
        ebitda = base_y5["subRevenue"] * cm * am + base_y5["adRevenue"] - base_y5["fixed"] - base_y5["payroll"]
        values.append(fmt_te(ebitda))
    sens_rows.append([f"{int(cm*100)}%"] + values)
add_table(doc,["Kundenvolumen / ARPA","80%","90%","100%","110%","120%"],sens_rows,[4.2,2.65,2.65,2.65,2.65,2.65],7.6)

# 18
chapter(doc, 18, "Kapitalbedarf, Mittelverwendung und Finanzierungsszenarien")
fund_p = DATA["scenarios"]["Pessimistisch"]["fundingNeed"]
fund_r = DATA["scenarios"]["Realistisch"]["fundingNeed"]
fund_o = DATA["scenarios"]["Optimistisch"]["fundingNeed"]
add_para(doc, f"Das Monatsmodell weist bei 2.000 € Startliquidität eine maximale rechnerische Finanzierungslücke von rund {fmt_num(fund_p/1000, 1)} T€ (pessimistisch), {fmt_num(fund_r/1000, 1)} T€ (realistisch) bzw. {fmt_num(fund_o/1000, 1)} T€ (optimistisch) aus. Dass das optimistische Szenario dennoch kurzzeitig negativ wird, folgt aus Gründungs-, Rechts- und Produktaufwand vor den jährlichen Vorauszahlungen. Ein reiner 2.000-€-Start ist daher unvorsichtig.")
add_table(doc, ["Szenario", "Struktur", "Mittel", "Eignung"], [
    ["A · Minimal Bootstrap", "2 T€ Stammkapital + 13 T€ Gesellschafterdarlehen", "15 T€", "Empfohlenes Minimum; enge Kostenkontrolle"],
    ["B · Robuster Bootstrap", "2 T€ Stammkapital + 23 T€ Darlehen/Fördermittel", "25 T€", "Bevorzugt; deckt Compliance, Versicherung und Puffer"],
    ["C · Wachstumsfinanzierung", "25–50 T€ Förderdarlehen/Bank nach Verträgen", "25–50 T€", "Erst bei 5+ zahlenden Kunden und klarer Pipeline"],
    ["D · Equity", "Angel/Seed", ">250 T€", "Derzeit nicht empfohlen; erst für belastbare Expansion"],
], [3.0, 6.0, 2.6, 6.0])
add_table(doc, ["Verwendung einer 20-T€-Reserve", "Betrag", "Begründung"], [
    ["Gründung, Verträge, Datenschutz, IP/Lizenzen", "4,5 T€", "Launch-Gates und anwaltliche Prüfung"],
    ["Versicherung, Buchhaltung, Verwaltung", "3,0 T€", "12–18 Monate Grundbetrieb"],
    ["Infrastruktur, Daten, Monitoring", "3,0 T€", "Last-/Lizenzpuffer"],
    ["Sales, Pilotintegration, Reisen", "4,0 T€", "Paid-Design-Partner-Akquise"],
    ["Contingency", "5,5 T€", "Ausfälle, Nachlizenzierung, längerer Sales-Zyklus"],
], [8.0, 2.7, 6.9])
add_para(doc, "Kapital wird nicht für breite Paid-Acquisition eingesetzt, bevor Conversion und Retention belegt sind. Gründerzeit bleibt der knappste wirtschaftliche Input; die Schattenkosten zeigen, dass „kostenlose“ Arbeit das Geschäftsmodell wirtschaftlich belastet.")

# 19
chapter(doc, 19, "Exit- beziehungsweise langfristige Wachstumsstrategie")
add_para(doc, "Primäres Ziel ist ein profitables, kapital-effizientes Softwareunternehmen. Ein Exit ist eine Option, kein Planungsersatz. Strategische Käufer könnten Sportdatenanbieter, Publisher-Softwaregruppen, Engagement-/CMS-Plattformen oder Sportmediennetzwerke sein, sofern das Unternehmen wiederkehrende Erlöse, Datenrechte, Prognosehistorie und Kundenbindung besitzt.")
add_table(doc, ["Pfad", "Voraussetzung", "Zeithorizont", "Priorität"], [
    ["Profitables Nischen-SaaS", "200–300 Kunden, geringe Churn, kleines Team", "3–5 Jahre", "Primär"],
    ["Europäische Expansion", "Deutschland-PMF, englisches Produkt, Rechte", "ab Jahr 3", "Sekundär"],
    ["US-Expansion", "US-Rechte, Support, Salespartner, Rechtsprüfung", "ab Jahr 4/5", "Optional"],
    ["Strategischer Verkauf", "ARR-Qualität, IP-Chain, Datenasset, NRR", "ab Jahr 5", "Opportunistisch"],
], [4.3, 7.3, 3.1, 3.0])
add_para(doc, "Exit-Readiness-KPIs: Anteil wiederkehrender Umsätze >80%, NRR ≥100%, Bruttomarge >70%, Vollkosten-CAC-Payback <12 Monate, geringe Kundenkonzentration, vollständige IP-Chain-of-Title und dokumentierte Modell-/Daten-Governance.")

# 20
chapter(doc, 20, "Abschließende Investitionsargumentation")
add_para(doc, "Das Vorhaben besitzt eine ungewöhnlich starke Kombination aus technischer Gründerqualität, vorhandener öffentlicher Resonanz und nahezu marktreifem Produkt. Gleichzeitig sind die entscheidenden kommerziellen Beweise noch offen. Für eine institutionelle Beteiligung fehlen zahlende Kunden, Kohorten-Retention, wiederholbare Akquisition, geklärte IP-/Datenrechte und ein hinreichender Team-Zeitcommitment.")
add_table(doc, ["Argument für", "Argument gegen", "Urteil"], [
    ["Technisch glaubwürdiges Founder-Team", "Nur ca. 20 Founder-Stunden/Woche gesamt", "Kapazität strikt priorisieren"],
    ["Belegte Nutzer-/Medienresonanz", "Eventabhängig; keine bezahlte B2B-Traktion", "Guter Lead-Magnet, kein PMF-Beweis"],
    ["Niedrige heutige Infrastrukturkosten", "Recht, Vertrieb, Support und Gründerzeit fehlen in Ist-Kosten", "Schattenrechnung beibehalten"],
    ["SaaS mit potenziell guter Marge", "ARPA/Churn/CAC unvalidiert", "Paid-Pilot-Phase zwingend"],
    ["Erklärbarkeit und Wissenschaft", "Prompts leicht kopierbar; Paperstatus offen", "Daten-/Benchmark-Moat aufbauen"],
], [6.2, 7.1, 4.7])
add_callout(doc, "Finales Votum", "Conditional Go für einen 90-tägigen, bootstrapped Markttest; No-Go für eine Equity-Runde zum jetzigen Zeitpunkt. Investierbarkeit entsteht erst nach 3–5 zahlenden Designpartnern, geklärten Rechten und einer reproduzierbaren Prognose-Scorecard.")

# IC check
doc.add_page_break()
doc.add_heading("Investment-Committee-Check und eingearbeitete Überarbeitung", level=1)
add_para(doc, "Ein skeptisches Investment Committee würde das Vorhaben aktuell nicht aufgrund der Medienreichweite finanzieren. Die folgende Prüfung wurde als Red-Team-Schritt durchgeführt; die daraus abgeleiteten Änderungen sind bereits in dieser finalen Fassung umgesetzt.")
add_table(doc, ["Kriterium", "Score / 5", "IC-Kritik", "Eingearbeitete Änderung"], [
    ["Problem/Kundenbedarf", "3,0", "Interesse, aber keine Zahlungsdaten", "Paid-Design-Partner als 90-Tage-Gate"],
    ["Produkt", "4,0", "nahezu marktreif, aber QA/Rechte offen", "Launch-Gates und vier konkrete Widgets"],
    ["Markt", "3,0", "TAM begrenzt und nur Proxy", "bottom-up Bandbreite statt Marktfloskel"],
    ["Traktion", "3,0", "Reichweite ≠ Retention", "Screenshots präzise ausgewertet; Likes nicht behauptet"],
    ["Geschäftsmodell", "2,5", "ARPA, Churn, RPM unbekannt", "drei Szenarien + Sensitivität"],
    ["GTM", "2,0", "Founder-led Sales noch nicht bewiesen", "30/60/90-Tage-Plan mit Kill-Kriterien"],
    ["Moat", "2,5", "Prompts sind kein Schutzwall", "Forecast Ledger, Kalibrierung, Integration"],
    ["Team", "4,0", "stark technisch, geringe Zeit/Commercial Gap", "Rollen-/Hiring-Gates"],
    ["Governance/Regulatorik", "2,0", "IP, Universität, Daten, 50/50 offen", "Gründungs- und Rechte-Checkliste"],
    ["Finanzen", "3,0", "2 T€ zu knapp; Gratisarbeit verzerrt", "15–25 T€ Reserve + Schattenkosten"],
], [4.0, 2.2, 6.0, 6.0], 7.2)
add_para(doc, "Gesamtscore: 2,9/5. Das entspricht einem attraktiven Experiment mit begrenztem Kapitaleinsatz, aber noch keinem investorenreifen Scale-up. Die finale Fassung korrigiert vier typische Überbewertungen: Medienreichweite wird nicht als Nachfrage verkauft; Ads sind sekundär; Prompts gelten nicht als Moat; und die Finanzierung berücksichtigt Liquidität und Gründerarbeit.")
add_heading("Entscheidungsregeln nach 90 Tagen", level=2)
add_bullets(doc, [
    "GO: mindestens 5 zahlende Kunden, davon 3 nach 60 Tagen aktiv; dokumentierte Prognose-Scorecard; Rechte und Verträge geklärt.",
    "REVISE: 2–4 zahlende Kunden oder starke Nutzung ohne Retention; Segment, Preis oder Einbindungsformat in einem weiteren 60-Tage-Test ändern.",
    "STOP/PIVOT: weniger als 2 zahlende Kunden trotz 30 qualifizierter Gespräche; keine weitere Sportarten- oder Automatisierungsausweitung.",
])

# Evidence appendix
doc.add_page_break()
doc.add_heading("Anhang A · Traction-Evidenz", level=1)
add_para(doc, "Die Screenshots sind Momentaufnahmen. Vor externen Gesprächen sollten Originalexporte, Property-ID, Zeitraum, Zeitzone, Filter und Plattformlinks in einem unveränderlichen Datenraum archiviert werden.")
evidence = doc.add_table(rows=1, cols=2)
evidence.alignment = WD_TABLE_ALIGNMENT.CENTER
evidence.autofit = False
left, right = evidence.rows[0].cells
for c in (left, right):
    set_cell_margins(c, 80, 80, 80, 80)
if Path("/Users/jonasschroder/Downloads/IMG_2240.PNG").exists():
    left.paragraphs[0].add_run().add_picture("/Users/jonasschroder/Downloads/IMG_2240.PNG", width=Cm(6.0))
    left.add_paragraph("Google Analytics: 13.669 aktive Nutzer; Zeitraum 11.06.–19.07.2026.", style="Small Note")
if Path("/Users/jonasschroder/Downloads/IMG_2238.jpg").exists():
    right.paragraphs[0].add_run().add_picture("/Users/jonasschroder/Downloads/IMG_2238.jpg", width=Cm(5.7))
    right.add_paragraph("Tagesschau-Screenshot: 7,7 Mio. Videoaufrufe; Motiv datiert 23.06.2026.", style="Small Note")
add_table(doc, ["Kennzahl", "Wert", "Zeitraum/Beleg", "Bewertung"], [
    ["Pageviews", "69.919", "11.06.–19.07.; IMG_2244/2243", "belegt im Screenshot"],
    ["Sessions", "36.066", "11.06.–19.07.; IMG_2242", "belegt im Screenshot"],
    ["Aktive Nutzer", "13.669", "11.06.–19.07.; IMG_2240/2241", "belegt im Screenshot"],
    ["Neue Nutzer", "13.596", "11.06.–19.07.; IMG_2240/2241", "belegt im Screenshot"],
    ["Ø Interaktionsdauer", "ca. 2:22 min", "IMG_2241", "belegt im Screenshot"],
    ["Tagesschau-Videoaufrufe", "7,7 Mio.", "IMG_2238/2239", "belegt im Screenshot"],
    ["Likes", "nicht belegt", "nicht in Bildern sichtbar", "nicht verwenden"],
], [4.0, 2.8, 6.3, 4.5])

# Advisor checklist
doc.add_page_break()
doc.add_heading("Anhang B · Fragenkatalog für Steuerberatung, Recht und Versicherungen", level=1)
doc.add_heading("Steuerberatung", level=2)
add_bullets(doc, [
    "UG oder GmbH: Einlage, Gesellschafterdarlehen, Gründungsaufwand und spätere Umwandlung?",
    "Tatsächlicher Ort der Geschäftsleitung München/Grünwald; Gewerbesteuerzerlegung und Substanzanforderungen?",
    "Umsatzsteuer bei B2B-Kunden in EU/USA, Reverse Charge, USt-ID, OSS-Relevanz und Leistungsort?",
    "E-Rechnungsfähigkeit, Kontenrahmen, Abgrenzung jährlicher Vorauszahlungen und Deferred Revenue?",
    "Steuerliche Behandlung von Softwareentwicklung, Forschung, Fördermitteln und IP-Übertragung?",
])
doc.add_heading("Rechtsberatung", level=2)
add_bullets(doc, [
    "LMU/Nebentätigkeit, Erfindungs-/Software-/Datenrechte und Publikationspflichten geklärt?",
    "TheSportsDB und weitere Daten-/Logo-/Markenrechte für kommerzielle Website und Kundenwidgets dokumentiert?",
    "B2B-AGB, AVV/DPA, SLA, Support, 12-Monatslaufzeit, automatische Verlängerung und wirksame Haftungsgrenzen?",
    "AI-Act-, DSGVO-, TDDDG-, UWG- und Werbe-/Cookie-Compliance für Plattform, Analytics und Outreach?",
    "Gesellschaftervertrag: 50/50-Deadlock, Vesting, Leaver, IP, Zeitcommitment und Universitätseintritt?",
])
doc.add_heading("Versicherungsmakler", level=2)
add_bullets(doc, [
    "IT-/Vermögensschadenhaftpflicht: Sind Medien-, IP-, Datenschutz-, Datenverlust- und Eigenschäden gedeckt?",
    "Cyber: Incident Response, Forensik, Betriebsunterbrechung, Ransomware, Cloud-/Dienstleisterausfall?",
    "Betriebshaftpflicht und weltweiter Geltungsbereich einschließlich USA; Ausschlüsse für Sportwetten/Prognosen?",
    "Welche Mindestkontrollen werden verlangt: MFA, Backups, Patchmanagement, Logging, Verträge, Notfallplan?",
    "Ab wann sind D&O/Managerhaftpflicht und Firmenrechtsschutz wirtschaftlich sinnvoll?",
])

# Sources
doc.add_page_break()
doc.add_heading("Anhang C · Quellenverzeichnis", level=1)
sources = [
    (1, "U.S. Small Business Administration – Write your business plan", "https://www.sba.gov/business-guide/plan-your-business/write-your-business-plan?showAll=true"),
    (2, "KfW – Gründung und Businessplan", "https://www.kfw.de/inlandsfoerderung/Unternehmen/Gr%C3%BCndung-und-Nachfolge/Gr%C3%BCndung/"),
    (3, "IAB Europe – AdEx Benchmark 2025", "https://iabeurope.eu/knowledge_hub/iab-europe-adex-benchmark-2025-report/"),
    (4, "BVDW/OVK – Digitaler Werbemarkt 2026", "https://www.bvdw.org/news-und-publikationen/ovk-prognose-digitaler-werbemarkt-waechst-auf-ueber-acht-milliarden-euro/"),
    (5, "IAB – U.S. Digital Ad Revenue 2025", "https://www.iab.com/news/digital-ad-revenue-climbs-to-nearly-300b-as-iab-celebrates-30-year-anniversary/"),
    (6, "Eurostat – Publishing activities statistics", "https://ec.europa.eu/eurostat/web/products-eurostat-news/w/ddn-20250502-1"),
    (7, "Eurostat – Online news consumption 2024", "https://ec.europa.eu/eurostat/web/products-eurostat-news/w/ddn-20250710-2"),
    (8, "BDZV – Reichweiten der Zeitungen", "https://www.bdzv.de/alle-themen/marktdaten/reichweiten-der-zeitungen"),
    (9, "Reuters Institute – Journalism, media and technology trends 2026", "https://reutersinstitute.politics.ox.ac.uk/journalism-media-and-technology-trends-and-predictions-2026"),
    (10, "BDZV/Highberg – Trends der Zeitungsbranche 2025", "https://www.bdzv.de/service/presse/pressemitteilungen/2025/trends-der-zeitungsbranche-2025"),
    (11, "Opta Analyst – Supercomputer methodology example", "https://theanalyst.com/articles/serie-a-predictions-2025-26-opta-supercomputer"),
    (12, "Sportmonks – Football Widgets Pricing", "https://www.sportmonks.com/football-api/football-widgets/"),
    (13, "API-Sports – Widgets documentation", "https://api-sports.io/documentation/widgets/v3"),
    (14, "ScoreAxis – Sports widgets", "https://www.scoreaxis.com/widgets/"),
    (15, "Google AdSense – Revenue share", "https://support.google.com/adsense/answer/180195?hl=de"),
    (16, "EU Commission – AI Act transparency guidelines", "https://digital-strategy.ec.europa.eu/en/news/commission-publishes-guidelines-transparency-obligations-providers-and-deployers-certain-ai-systems"),
    (17, "EU Commission – AI regulatory framework", "https://digital-strategy.ec.europa.eu/en/policies/regulatory-framework-ai"),
    (18, "LMU – Inventions, patents and exploitation rights", "https://www.lmu.de/en/research/research-transfer/inventions-patents-and-exploitation-rights/"),
    (19, "Gesetze im Internet – GmbHG §5/§5a", "https://www.gesetze-im-internet.de/gmbhg/BJNR004770892.html"),
    (20, "Bundesfinanzministerium – FAQ E-Rechnung", "https://www.bundesfinanzministerium.de/Content/DE/FAQ/e-rechnung.html"),
    (21, "IHK Köln – Versicherungen für Unternehmen", "https://www.ihk.de/koeln/hauptnavigation/beratung-und-services/versicherungen-fuer-unternehmen-5141790"),
    (22, "Stadt München – Gewerbesteuerhebesatz", "https://stadt.muenchen.de/infos/hebesaetze-gewerbesteuer-grundsteuer.html"),
    (23, "KStG §23", "https://www.gesetze-im-internet.de/kstg_1977/__23.html"),
    (24, "SolZG §4", "https://www.gesetze-im-internet.de/solzg_1995/__4.html"),
    (25, "GewStG – Steuermesszahl", "https://www.gesetze-im-internet.de/gewstg/BJNR009790936.html"),
    (26, "LLM SoccerArena – About", "https://www.llm-soccerarena.com/about"),
    (27, "BR – Bericht zur KI-Sportprognose", "https://www.br.de/nachrichten/wissen/neue-studie-wie-gut-prognostiziert-die-ki-fussball-ergebnisse%2CVMyclTs"),
]
for idx, label, url in sources:
    p = doc.add_paragraph(style="Small Note")
    p.paragraph_format.left_indent = Cm(0.25)
    p.paragraph_format.first_line_indent = Cm(-0.25)
    p.add_run(f"[Q{idx}] {label}: ").bold = True
    add_hyperlink(p, url, url)

doc.add_heading("Noch einzureichende Nachweise", level=2)
add_bullets(doc, [
    "Paper-URL/DOI/arXiv- oder Archiv-Identifier, Version, Einreichungsstatus und freigegebene Benchmark-Ergebnisse.",
    "Originale Google-Analytics-Exporte und Plattformlinks zur Tagesschau-/BR-Berichterstattung.",
    "Interviewliste mit Datum, Rolle, Problem, Budget, Kaufprozess und anonymisiertem Zitat.",
    "TheSportsDB-/API-Vertrag sowie Logo-/Markenfreigaben für kommerzielle Nutzung.",
    "Universitätsfreigabe, Nebentätigkeitsanzeige und dokumentierte IP-Chain-of-Title.",
])

doc.add_heading("Finale Freigabeliste vor Gründung", level=2)
add_numbered(doc, [
    "Unternehmensname, Domain und Markenrecherche entscheiden.",
    "LMU-/IP-/Nebentätigkeitsklärung schriftlich abschließen.",
    "Daten-, Logo- und API-Lizenzen kommerziell freigeben lassen.",
    "UG-Satzung und 50/50-Gesellschaftervertrag notariell vorbereiten.",
    "15–25 T€ Liquiditätsreserve verbindlich dokumentieren.",
    "AGB, Datenschutz, DPA/AVV, Impressum und KI-Hinweise prüfen.",
    "IT-/Vermögensschaden- und Cyberdeckung anbieten lassen.",
    "Vier Widgets mit Monitoring, Scorecard und Rollback abnehmen.",
    "30 Zielpublisher priorisieren und Paid-Pilot-Angebot standardisieren.",
    "90-Tage-Go/Revise/Stop-Review im Gründerkalender terminieren.",
])

doc.add_paragraph()
add_callout(doc, "Schlussfolgerung", "Das Projekt ist wirtschaftlich plausibel, wenn es als fokussiertes Publisher-SaaS geführt wird. Es ist noch nicht bewiesen, dass es skalierbar ist. Der nächste Unternehmenswert entsteht nicht durch weitere Features, sondern durch zahlende Kunden, belastbare Retention, messbare Prognosegüte und geklärte Rechte.")

# Metadata and save
doc.core_properties.title = "AI Sports Prediction – Businessplan 2026–2031"
doc.core_properties.subject = "Finaler, investorenfähiger Businessplan für interne Planung, Steuerberatung und Versicherungen"
doc.core_properties.author = "Jonas Schröder und Jonas Schweisthal"
doc.core_properties.keywords = "Businessplan, AI, Sports Prediction, Publisher Widgets, SaaS, Finanzplanung"
target = OUT / "AI_Sports_Prediction_Businessplan_2026.docx"
doc.save(target)
print(target)
